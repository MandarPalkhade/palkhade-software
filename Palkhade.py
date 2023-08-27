import tkinter as tk
from tkinter import ttk, messagebox
import pandas as pd
import datetime
import docx
import os
from docx.shared import Inches
#Mandar


class DocumentGenerator:
    def __init__(self, master):
        self.master = master
        self.master.title("Document Generator")

        # Load doctor names from Excel
        self.doctor_names = self.load_doctor_names_from_excel("D:\\PALKHADE LAB REPORTS\\doctors.xlsx")

        # Create fonts
        font_bold = ("Arial", 12, "bold")
        font = ("Arial", 12)

        # Create labels and entry boxes for patient name, doctor name, and date
        ttk.Label(master, text="Patient Name: ", font=font_bold).grid(row=0, column=0, padx=5, pady=5)
        self.patient_name_entry = ttk.Entry(master, width=30, font=font)
        self.patient_name_entry.grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(master, text="Doctor Name: ", font=font_bold).grid(row=1, column=0, padx=5, pady=5)
        self.doctor_name_entry = ttk.Entry(master, width=30, font=font)
        self.doctor_name_entry.grid(row=1, column=1, padx=5, pady=5)

        # Add a dropdown list for selecting the doctor's name
        ttk.Label(master, text="Or select from the list below:", font=font_bold).grid(row=2, column=0, padx=5, pady=5)
        self.doctor_name_selected = tk.StringVar()
        self.doctor_name_selected.set(self.doctor_names[0])
        self.doctor_name_dropdown = ttk.Combobox(master, values=self.doctor_names, textvariable=self.doctor_name_selected, state="readonly", font=font)
        self.doctor_name_dropdown.grid(row=2, column=1, padx=5, pady=5)

        # Add a label and entry box for adding a new doctor's name
        ttk.Label(master, text="New Doctor Name: ", font=font_bold).grid(row=3, column=0, padx=5, pady=5)
        self.new_doctor_name_entry = ttk.Entry(master, width=30, font=font)
        self.new_doctor_name_entry.grid(row=3, column=1, padx=5, pady=5)

        ttk.Button(master, text="Add Doctor", command=self.add_doctor_name).grid(row=4, column=0, columnspan=2, padx=5, pady=5)

        ttk.Label(master, text="Gender: ", font=font_bold).grid(row=5, column=0, padx=5, pady=5)
        self.gender_selected = tk.StringVar()
        self.gender_selected.set("Male")
        self.gender_dropdown = ttk.Combobox(master, values=["Male", "Female"], textvariable=self.gender_selected, state="readonly", font=font)
        self.gender_dropdown.grid(row=5, column=1, padx=5, pady=5)

        ttk.Label(master, text="Date: ", font=font_bold).grid(row=6, column=0, padx=5, pady=5)
        self.date_label = ttk.Label(master, text=datetime.date.today().strftime("%d/%m/%Y"), font=font)
        self.date_label.grid(row=6, column=1, padx=5, pady=5)

        # Create button to generate document
        ttk.Button(master, text="Generate Document", command=self.generate_document).grid(row=7, column=0, columnspan=2, padx=5, pady=5)

        # Bind the "Enter" key to the "Generate Document" button
        master.bind("<Return>", lambda event: self.generate_document())

        # Configure grid to expand rows and columns
        master.grid_rowconfigure(8, weight=1)
        master.grid_columnconfigure(0, weight=1)
        master.grid_columnconfigure(1, weight=1)

        # Bind the keypress event to the doctor_name_entry field
        self.doctor_name_entry.bind("<KeyRelease>", self.auto_suggest_doctor_name)

        # Create a listbox for displaying auto-suggestions
        self.suggestion_listbox = tk.Listbox(master, font=font)
        self.suggestion_listbox.grid(row=1, column=2, padx=5, pady=5)
        self.suggestion_listbox.bind("<<ListboxSelect>>", self.select_suggested_doctor_name)

    def auto_suggest_doctor_name(self, event):
        # Get the typed text from the doctor_name_entry field
        entered_text = self.doctor_name_entry.get()

        # Clear the suggestion listbox
        self.suggestion_listbox.delete(0, tk.END)

        # Filter doctor names that match the entered text
        matching_doctors = [doctor for doctor in self.doctor_names if doctor.lower().startswith(entered_text.lower())]

        # Add the matching doctors to the suggestion listbox
        for doctor in matching_doctors:
            self.suggestion_listbox.insert(tk.END, doctor)

    def select_suggested_doctor_name(self, event):
        # Get the selected doctor name from the suggestion listbox
        selected_doctor = self.suggestion_listbox.get(tk.ACTIVE)

        # Set the selected doctor name in the doctor_name_entry field
        self.doctor_name_entry.delete(0, tk.END)
        self.doctor_name_entry.insert(tk.END, selected_doctor)

    def generate_document(self):
        # Get the patient name, doctor name, date, and gender entered
        patient_name = self.patient_name_entry.get()
        doctor_name = self.doctor_name_entry.get() or self.doctor_name_selected.get()
        gender = self.gender_selected.get()
        date = self.date_label["text"]

        # Create a new document
        document = docx.Document()

        # Set the page margins
        section = document.sections[0]
        section.top_margin = Inches(1.97)
        section.bottom_margin = Inches(1.18)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

        # Add the doctor name, patient name, and date to the document
        paragraph = document.add_paragraph()
        paragraph.add_run(f"Patient Name: {patient_name}\t\t\t\t\t\tDate: {date}").font.size = docx.shared.Pt(14)
        paragraph = document.add_paragraph()
        paragraph.add_run(f"Ref. by Dr: {doctor_name}\t\t\t\t\tGender: {gender}").font.size = docx.shared.Pt(14)

        # Save the document in D:\JAVA folder with a unique filename based on the current date and time
        folder_path = r"D:\\PALKHADE LAB REPORTS"
        filename = os.path.join(folder_path, f"{patient_name}_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx")
        document.save(filename)

        # Open the document with the default application
        os.startfile(filename)

        # Show a message box confirming that the document was generated and saved
        messagebox.showinfo("Document Generated", f"Your document was generated and saved as {filename}.")

    def load_doctor_names_from_excel(self, filename):
        # Load doctor names from an Excel file using pandas
        df = pd.read_excel(filename)
        return df["Doctor Name"].tolist()

    def add_doctor_name(self):
        # Get the new doctor's name entered
        new_doctor_name = self.new_doctor_name_entry.get()

        # Append the new doctor's name to the Excel file
        filename = "D:\\PALKHADE LAB REPORTS\\doctors.xlsx"
        df = pd.read_excel(filename)
        df = pd.concat([df, pd.DataFrame({"Doctor Name": [new_doctor_name]})])
        df.to_excel(filename, index=False)

        # Update the doctor names list
        self.doctor_names = self.load_doctor_names_from_excel(filename)
        self.doctor_name_dropdown["values"] = self.doctor_names


root = tk.Tk()
root.geometry("600x450")  # Set the size of the window (width x height)
document_generator = DocumentGenerator(root)
root.mainloop()
